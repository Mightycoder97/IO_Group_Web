import sys
import json
from docx import Document
import re

def replace_text_in_paragraph(paragraph, replacements):
    # docx splits text into runs which can break placeholders
    # The safest way is to rebuild the paragraph text if a placeholder is found
    # but that loses run-level formatting.
    # We will try a simple replace on the entire paragraph text for robustness.
    original_text = paragraph.text
    new_text = original_text
    
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, str(value))
            
    if new_text != original_text:
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        # Append the new text to the first run to keep block formatting
        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

def replace_text_in_doc(doc, replacements):
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, replacements)

def main():
    if len(sys.argv) < 4:
        print(json.dumps({"success": False, "error": "Missing arguments"}))
        sys.exit(1)

    template_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        data = json.loads(sys.argv[3])
    except json.JSONDecodeError:
        print(json.dumps({"success": False, "error": "Invalid JSON data"}))
        sys.exit(1)

    try:
        doc = Document(template_path)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"Error reading template: {e}"}))
        sys.exit(1)

    # Flatten the JSON to generic placeholders: [cliente_nombre], [empresa_razon_social], etc.
    replacements = {}
    for section, values in data.items():
        if isinstance(values, dict):
            for k, v in values.items():
                replacements[f"[{section}_{k}]"] = str(v)
        else:
            replacements[f"[{section}]"] = str(values)
            
    # Also support common possible tags
    if 'empresa' in data:
        replacements['<<RAZON_SOCIAL>>'] = str(data['empresa'].get('razon_social', ''))
        replacements['<<RUC>>'] = str(data['empresa'].get('ruc', ''))
        replacements['<<DIRECCION>>'] = str(data['empresa'].get('direccion_fiscal', ''))
    
    if 'cliente' in data:
        replacements['<<REPRESENTANTE>>'] = str(data['cliente'].get('nombre', ''))
        replacements['<<DNI>>'] = str(data['cliente'].get('dni', ''))
        
    if 'sede' in data:
        replacements['<<SEDE_DIRECCION>>'] = str(data['sede'].get('direccion', ''))
        replacements['<<SEDE_DISTRITO>>'] = str(data['sede'].get('distrito', ''))
        
    # Generate pricing logic dynamically
    if 'contrato' in data:
        tipo_tarifa = data['contrato'].get('tipo_tarifa', 'por_servicio')
        tarifa = data['contrato'].get('tarifa', '0')
        peso_limite = data['contrato'].get('peso_limite_kg', '')
        tarifa_adicional = data['contrato'].get('tarifa_adicional_kg', '')
        
        if tipo_tarifa == 'por_kg':
            clausula = f"El servicio se establece bajo la modalidad 'Por Kilo Recogido'. El costo base establecido cubrirá un límite de {peso_limite} Kg. Por cada kilo excedente, se facturará un adicional de S/. {tarifa_adicional}."
            replacements['[clausula_exceso_peso]'] = clausula
            replacements['<<CLAUSULA_EXCESO>>'] = clausula
        else:
            clausula = f"El servicio se establece bajo la modalidad 'Por Recojo' con una tarifa plana por atención."
            replacements['[clausula_exceso_peso]'] = clausula
            replacements['<<CLAUSULA_EXCESO>>'] = clausula
            
    replace_text_in_doc(doc, replacements)

    try:
        doc.save(output_path)
        print(json.dumps({
            "success": True, 
            "message": "Contract generated", 
            "file": output_path,
            "placeholders_used": list(replacements.keys())
        }))
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
