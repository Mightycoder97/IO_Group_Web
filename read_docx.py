import docx
import re
import sys

def extract_placeholders(filename):
    try:
        doc = docx.Document(filename)
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        # Look for common placeholder patterns, e.g., <<VAR>>, [VAR], {VAR}, {{VAR}}
        patterns = [
            r'<<([^>]+)>>',
            r'\[([^\]]+)\]',
            r'\{([^\}]+)\}',
            r'\{\{([^\}]+)\}\}'
        ]
        
        found = set()
        for p in patterns:
            matches = re.findall(p, text)
            for m in matches:
                found.add(m.strip())
                
        print("Found potential placeholders:")
        for f in sorted(list(found)):
            print(f"- {f}")
            
        # Also print first 1000 chars of text to inspect manually if needed
        print("\n\nFirst 500 chars of text:")
        print(text[:500])
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_placeholders(sys.argv[1])
